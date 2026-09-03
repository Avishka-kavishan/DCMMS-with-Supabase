import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def create_full_document():
    doc = Document()
    
    # Page Setup - Standard A4 with 1-inch margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
        # Header / Footer setup
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hrun = hp.add_run("DCMMS — System Administrator Handover & Maintenance Guide")
        hrun.font.name = "Arial"
        hrun.font.size = Pt(8.5)
        hrun.font.color.rgb = RGBColor(148, 163, 184)
        
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        frun = fp.add_run("Ministry of Education, Sri Lanka — Discipline Branch | Confidential Government Asset")
        frun.font.name = "Arial"
        frun.font.size = Pt(8)
        frun.font.color.rgb = RGBColor(148, 163, 184)

    def set_cell_background(cell, fill_hex):
        tcPr = cell._element.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
        tcPr.append(shd)

    def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
        tcPr = cell._element.get_or_add_tcPr()
        tcMar = parse_xml(f'''
            <w:tcMar {nsdecls("w")}>
                <w:top w:w="{top}" w:type="dxa"/>
                <w:bottom w:w="{bottom}" w:type="dxa"/>
                <w:left w:w="{left}" w:type="dxa"/>
                <w:right w:w="{right}" w:type="dxa"/>
            </w:tcMar>
        ''')
        tcPr.append(tcMar)

    def set_table_borders(table, color="CBD5E1"):
        tblPr = table._element.xpath('w:tblPr')
        if tblPr:
            borders = parse_xml(f'''
                <w:tblBorders {nsdecls("w")}>
                    <w:top w:val="single" w:sz="6" w:space="0" w:color="{color}"/>
                    <w:bottom w:val="single" w:sz="6" w:space="0" w:color="{color}"/>
                    <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{color}"/>
                    <w:insideV w:val="none"/>
                    <w:left w:val="none"/>
                    <w:right w:val="none"/>
                </w:tblBorders>
            ''')
            tblPr[0].append(borders)

    def add_title(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(36)
        p.paragraph_format.space_after = Pt(12)
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(24)
        run.bold = True
        run.font.color.rgb = RGBColor(10, 37, 64) # Navy

    def add_subtitle(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(24)
        run = p.add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(71, 85, 105)

    def add_h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(20)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(15)
        run.bold = True
        run.font.color.rgb = RGBColor(10, 37, 64)

    def add_h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(12.5)
        run.bold = True
        run.font.color.rgb = RGBColor(30, 41, 59)

    def add_h3(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(10.5)
        run.bold = True
        run.font.color.rgb = RGBColor(51, 65, 85)

    def add_p(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(30, 41, 59)
        return p

    def add_bullet(text, level=0):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.1
        run = p.add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(30, 41, 59)
        return p

    def add_callout(title, text, box_type="NOTE"):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit = False
        
        color_map = {
            "NOTE": ("F0F9FF", "0284C7"),
            "IMPORTANT": ("F0FDF4", "16A34A"),
            "WARNING": ("FFFBEB", "D97706"),
            "CAUTION": ("FEF2F2", "DC2626")
        }
        bg_hex, border_hex = color_map.get(box_type, ("F8FAFC", "475569"))
        
        cell = tbl.cell(0, 0)
        cell.width = Inches(6.5)
        set_cell_background(cell, bg_hex)
        set_cell_margins(cell, top=140, bottom=140, left=180, right=160)
        
        tcPr = cell._element.get_or_add_tcPr()
        borders = parse_xml(f'''
            <w:tcBorders {nsdecls("w")}>
                <w:left w:val="single" w:sz="24" w:space="0" w:color="{border_hex}"/>
                <w:top w:val="none"/>
                <w:right w:val="none"/>
                <w:bottom w:val="none"/>
            </w:tcBorders>
        ''')
        tcPr.append(borders)
        
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run_title = p.add_run(f"[{box_type}] {title}\n")
        run_title.bold = True
        run_title.font.name = "Arial"
        run_title.font.size = Pt(9.5)
        r, g, b = tuple(int(border_hex[i:i+2], 16) for i in (0, 2, 4))
        run_title.font.color.rgb = RGBColor(r, g, b)
        
        run_text = p.add_run(text)
        run_text.font.name = "Calibri"
        run_text.font.size = Pt(9)
        run_text.font.color.rgb = RGBColor(51, 65, 85)
        
        doc.add_paragraph().paragraph_format.space_after = Pt(3)

    def add_code_block(code_text):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit = False
        
        cell = tbl.cell(0, 0)
        cell.width = Inches(6.5)
        set_cell_background(cell, "F8FAFC")
        set_cell_margins(cell, top=100, bottom=100, left=140, right=140)
        
        tcPr = cell._element.get_or_add_tcPr()
        borders = parse_xml(f'''
            <w:tcBorders {nsdecls("w")}>
                <w:left w:val="single" w:sz="16" w:space="0" w:color="0A2540"/>
                <w:top w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>
                <w:right w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>
                <w:bottom w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>
            </w:tcBorders>
        ''')
        tcPr.append(borders)
        
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(code_text.strip())
        run.font.name = "Consolas"
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(30, 41, 59)
        doc.add_paragraph().paragraph_format.space_after = Pt(3)

    def add_custom_table(col_widths, headers, data):
        tbl = doc.add_table(rows=1, cols=len(headers))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit = False
        set_table_borders(tbl)
        
        hdr_cells = tbl.rows[0].cells
        for i, title in enumerate(headers):
            hdr_cells[i].text = title
            set_cell_background(hdr_cells[i], "0A2540")
            set_cell_margins(hdr_cells[i], top=100, bottom=100, left=100, right=100)
            p = hdr_cells[i].paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            for run in p.runs:
                run.font.bold = True
                run.font.name = "Arial"
                run.font.size = Pt(8.5)
                run.font.color.rgb = RGBColor(255, 255, 255)
                
        for row_idx, row_data in enumerate(data):
            row_cells = tbl.add_row().cells
            bg_color = "F8FAFC" if row_idx % 2 == 1 else "FFFFFF"
            for col_idx, text_val in enumerate(row_data):
                row_cells[col_idx].text = str(text_val)
                set_cell_background(row_cells[col_idx], bg_color)
                set_cell_margins(row_cells[col_idx], top=70, bottom=70, left=90, right=90)
                p = row_cells[col_idx].paragraphs[0]
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(8.5)
                    run.font.color.rgb = RGBColor(30, 41, 59)
                    
        for row in tbl.rows:
            for idx, width in enumerate(col_widths):
                row.cells[idx].width = width
                
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ----------------------------------------------------
    # COVER PAGE
    # ----------------------------------------------------
    add_title("Discipline Case Management and Monitoring System (DCMMS)")
    add_subtitle("System Administrator Handover, Maintenance, Upgrade and Release Management Guide\nMinistry of Education, Sri Lanka — Discipline Branch")
    
    add_p("Document Reference: DOC-DCMMS-SYSADMIN-GUIDE-V1.0\nClassification: Restricted Technical Asset\nPrimary System Architecture: Next.js 16 (React 19) / PostgreSQL 16 / Prisma ORM / PM2\nDeployment Scope: Local Server / Air-Gapped Government Intranet")
    doc.add_page_break()

    # ----------------------------------------------------
    # DOCUMENT CONTROL & APPROVAL
    # ----------------------------------------------------
    add_h1("Document Control & Approval")
    
    add_h2("Document Metadata")
    add_custom_table(
        [Inches(2.0), Inches(4.5)],
        ["Property", "Specification Details"],
        [
            ["System Title", "Discipline Case Management and Monitoring System (DCMMS)"],
            ["Document Title", "DCMMS System Administrator Handover, Maintenance & Release Guide"],
            ["Target Environment", "Ministry of Education Intranet / Local Server Computer"],
            ["Current Release Version", "Version 1.0.0 (Production Stable)"],
            ["Designated Custodian", "Lead System Administrator / Technical Officer, Ministry of Education"],
            ["Original Codebase Path", "e:\\DCMMS-with-Supabase (Local Git Repository)"]
        ]
    )

    add_h2("Revision History")
    add_custom_table(
        [Inches(0.8), Inches(1.1), Inches(1.8), Inches(2.8)],
        ["Version", "Date", "Author / Designation", "Summary of Changes"],
        [
            ["0.1.0", "2026-06-15", "Lead Developer", "Initial architecture draft and environment configuration."],
            ["0.9.0", "2026-08-10", "Full-Stack Intern", "Unified PostgreSQL 18-table schema and automation scripts."],
            ["1.0.0", "2026-09-01", "Software Engineering Team", "Full institutional handover guide: operations, migrations, and DRP."]
        ]
    )

    add_h2("Approval and Sign-off Table")
    add_custom_table(
        [Inches(1.8), Inches(2.2), Inches(1.5), Inches(1.0)],
        ["Role / Designation", "Officer Name", "Signature", "Date"],
        [
            ["Outgoing Lead Developer", "Avishka Kavishan (Full-Stack Intern)", "________________", "2026-09-01"],
            ["Incoming System Administrator", "[TO BE COMPLETED BY FUTURE SYSADMIN]", "________________", "____/____/2026"],
            ["Discipline Branch Head", "[TO BE CONFIRMED - Senior Asst. Secretary]", "________________", "____/____/2026"],
            ["Director of ICT (MOE)", "[TO BE CONFIRMED - Director of ICT]", "________________", "____/____/2026"]
        ]
    )

    doc.add_page_break()

    # ----------------------------------------------------
    # SECTION 1: DOCUMENT PURPOSE & OPERATIONAL SCOPE
    # ----------------------------------------------------
    add_h1("1. Document Purpose & Operational Scope")
    add_p("This document serves as the authoritative, institutional technical handover manual for the Discipline Case Management and Monitoring System (DCMMS) deployed at the Ministry of Education (MOE), Sri Lanka. The primary purpose is to guarantee operational continuity, system stability, architectural integrity, and independent maintainability after the departure of the original development team and software engineering interns.")
    add_p("A new technical officer, systems administrator, or database administrator taking custody of DCMMS must be able to deploy, configure, secure, troubleshoot, backup, migrate, and extend the system without requiring real-time consultation with prior development personnel.")

    add_h2("1.1 Operational Responsibility Matrix (RACI)")
    add_custom_table(
        [Inches(2.5), Inches(1.0), Inches(1.0), Inches(1.0), Inches(1.0)],
        ["Operational Area / Task", "Sys Admin", "Developer", "Branch Admin", "Officers"],
        [
            ["Server OS & Hardware Health", "R / A", "C", "I", "I"],
            ["PostgreSQL Database & Backups", "R / A", "C", "I", "I"],
            ["User Account Creation & Role Assignment", "R / A", "I", "C", "I"],
            ["Code Bug Fixes & New Releases", "C / A", "R", "C", "I"],
            ["Prisma Schema Migrations", "A", "R", "I", "I"],
            ["Case Filing & Disciplinary Minutes", "I", "I", "R / A", "R"],
            ["Daily Mail Registration & PDF Entry", "I", "I", "I", "R / A"],
            ["Excel Export Master Password Update", "R / A", "I", "C", "I"],
            ["Security Audit Log Forensic Review", "R / A", "C", "I", "I"],
            ["Emergency Disaster Recovery Execution", "R / A", "C", "I", "I"]
        ]
    )

    add_callout("Strict Maintenance Mandate", "No production deployment, database schema modification, environment change, or architectural refactor is considered complete until this guide, its appendices, and the associated change logs are updated and formally approved.", "IMPORTANT")

    # ----------------------------------------------------
    # SECTION 2: SYSTEM OVERVIEW & BUSINESS CONTEXT
    # ----------------------------------------------------
    add_h1("2. System Overview & Business Context")
    add_p("The Discipline Case Management and Monitoring System (DCMMS) is an enterprise web application custom-built for the Ministry of Education (MOE), Sri Lanka. It digitizes the end-to-end lifecycle of disciplinary complaints, investigations, and formal proceedings regarding educational officers, principals, teachers, and non-academic staff across 15 provincial/zonal administrative categories.")

    add_h2("2.1 Solved Organizational Problems")
    add_bullet("Elimination of untracked physical mail through centralized Daily Mail letter logging and automated PDF complaint attachment.")
    add_bullet("Automated statutory inquiry deadline tracking (+30 / +60 days) with proactive visual warning badges for approaching (<=14 days) and overdue investigations.")
    add_bullet("Algorithmic Inquirer Conflict-of-Interest prevention, verifying inquirer employment history and children's school affiliations against the accused officer's school.")
    add_bullet("Controlled Excel report generation protected by a master administrative password and immutable PostgreSQL audit trails.")

    add_h2("2.2 Demarcation: Computerized vs. Manual / Offline Processes")
    add_custom_table(
        [Inches(3.2), Inches(3.3)],
        ["Computerized Inside DCMMS", "Manual / Offline / Physical Government Process"],
        [
            ["Recording incoming mail and letter metadata", "Physical sorting, opening, and stamping of official postal envelopes."],
            ["Uploading and viewing scanned PDF complaint letters", "Physical file storage in secure Ministry discipline record vaults."],
            ["Committee member registry and conflict validation", "Physical interviews, witness depositions, and tribunal hearing sessions."],
            ["Setting and tracking appointment and due dates", "Physical drafting and manual ink-signing of official inquiry letters."],
            ["Extension request approval workflow", "Physical formal correspondence between Inquiry Officers and Ministry Secretary."],
            ["Disciplinary order recording and audit history", "Enforcement of Public Service Commission (PSC) salary cuts or interdictions."]
        ]
    )

    # ----------------------------------------------------
    # SECTION 3: SYSTEM ARCHITECTURE & INFRASTRUCTURE DESIGN
    # ----------------------------------------------------
    add_h1("3. System Architecture & Infrastructure Design")
    add_p("DCMMS is engineered as an air-gapped / intranet-hosted 3-tier system within the Ministry of Education Government Local Area Network (LAN):")
    add_bullet("Tier 1 - Client Workstations: Modern web browsers (Microsoft Edge, Google Chrome, Mozilla Firefox) accessing the intranet URL.")
    add_bullet("Tier 2 - Application & Reverse Proxy: IIS 10 / Nginx terminating port 80/443 and proxying to Next.js 16 running on Node.js 20 LTS via PM2 Cluster Mode (Port 3000).")
    add_bullet("Tier 3 - Data & Storage Layer: PostgreSQL 16 database running on port 5432 and local disk storage at public/uploads/documents/ for PDF documents.")

    add_h2("3.1 Component Failure & Troubleshooting Reference")
    add_custom_table(
        [Inches(1.5), Inches(1.5), Inches(1.5), Inches(2.0)],
        ["Component", "Configuration", "Failure Symptom", "Troubleshooting Action"],
        [
            ["Reverse Proxy (IIS/Nginx)", "web.config / nginx.conf", "HTTP 502 Bad Gateway", "Check PM2 port 3000 status: netstat -ano | findstr 3000"],
            ["Next.js Web Engine", "ecosystem.config.js", "App crash / 500 error", "Check PM2 logs: pm2 logs dmms-app; Restart: pm2 restart dmms-app"],
            ["Prisma ORM Client", "lib/prisma.ts", "DB connection exceptions", "Run npx prisma generate to synchronize TypeScript bindings"],
            ["PostgreSQL 16 Engine", "postgresql.conf / pg_hba.conf", "ECONNREFUSED 5432", "Check Windows Service postgresql-x64-16; verify disk space"],
            ["PDF Document Storage", "public/uploads/documents/", "Upload 500 / 404 links", "Check Windows NTFS permissions on public/uploads for IIS/PM2 user"]
        ]
    )

    # ----------------------------------------------------
    # SECTION 4: TECHNOLOGY STACK
    # ----------------------------------------------------
    add_h1("4. Technology Stack & Dependency Specifications")
    add_custom_table(
        [Inches(1.8), Inches(1.1), Inches(1.4), Inches(2.2)],
        ["Technology / Package", "Version", "Scope / Layer", "Operational Purpose"],
        [
            ["Next.js", "16.1.6", "Application Framework", "Full-stack framework (App router, Server Actions, SSR, API routes)."],
            ["React / React DOM", "19.2.3", "UI Engine", "Component rendering, reactive UI state, and client hydration."],
            ["TypeScript", "^5.0.0", "Programming Language", "Strict static typing across models, actions, and API payloads."],
            ["Tailwind CSS", "^4.0.0", "Styling Framework", "Design tokens, utility CSS, and responsive layout engine."],
            ["PostgreSQL", "16.x LTS", "Relational Database", "ACID-compliant storage for all 18 system tables and audit logs."],
            ["Prisma Client & CLI", "^5.22.0", "ORM & Migrations", "Type-safe database abstraction and schema push tooling."],
            ["pg / @types/pg", "^8.13.1", "Node Database Driver", "Underlying connection pooling driver for PostgreSQL."],
            ["PM2", "^5.3.x", "Process Manager", "Production clustering, CPU load balancing, and auto-restart."],
            ["Recharts", "^3.9.1", "Data Visualization", "Analytics charts, case distribution graphs, and session load."],
            ["Lucide React", "^1.23.0", "Iconography", "Standardized government UI icons."],
            ["i18next / react-i18next", "^26.3 / ^17.0", "Localization", "Sinhala and English bilingual interface dictionary engine."],
            ["Node.js", "v20.x LTS", "Runtime", "Server-side JavaScript runtime (Node v20.14.0+ recommended)."]
        ]
    )

    # ----------------------------------------------------
    # SECTION 5: SOURCE CODE STRUCTURE
    # ----------------------------------------------------
    add_h1("5. Source Code Repository Structure")
    add_p("The codebase is organized under a modular Next.js App Router architecture:")
    add_bullet("app/admin: Discipline Branch Administrator dashboard, case endorsements, and extension approvals.")
    add_bullet("app/daily-mail: Daily Mail officer dashboard for registering incoming postal letters and attaching PDF scans.")
    add_bullet("app/subject: Subject Officer workspace for case filing, stage progression, and disciplinary minutes.")
    add_bullet("app/investigation: Investigation Branch Administrator workspace for committee registry and appointment tracking.")
    add_bullet("app/system-admin: System Administrator dashboard for audit logs, active session management, user accounts, and master Excel export passwords.")
    add_bullet("app/api/upload: Dedicated endpoint enforcing 25MB limits, MIME validation, and unique disk naming.")
    add_bullet("lib/db-actions.ts: Core server actions containing 100% of Prisma PostgreSQL database operations (160KB+).")
    add_bullet("lib/security.ts: Security monitoring module managing session duration, remote forced logout, and audit trail inserts.")
    add_bullet("lib/subject-types.ts: Master dictionary of the 15 educational subject categories.")
    add_bullet("prisma/schema.prisma: Single source of truth database schema defining 18 unified tables.")
    add_bullet("scripts/backup_postgres.ps1: Production-ready PowerShell database backup automation script.")

    # ----------------------------------------------------
    # SECTION 6: ENVIRONMENT SETUP & INSTALLATION GUIDE
    # ----------------------------------------------------
    add_h1("6. Environment Setup & Installation Guide")
    add_p("Follow this step-by-step procedure to provision DCMMS on a fresh Windows Server or development workstation:")

    add_h2("Step 1: Install Node.js, PostgreSQL 16, and Git")
    add_code_block("node -v    # Expected: v20.x.x\nnpm -v     # Expected: 10.x.x\npsql --version # Expected: psql (PostgreSQL) 16.x")

    add_h2("Step 2: Initialize PostgreSQL Database Instance")
    add_code_block("CREATE DATABASE dmms_db;\nCREATE USER db_user WITH PASSWORD 'StrongGovernmentPass2026!';\nGRANT ALL PRIVILEGES ON DATABASE dmms_db TO db_user;\nALTER DATABASE dmms_db OWNER TO db_user;\n\\c dmms_db\nGRANT ALL ON SCHEMA public TO db_user;\nGRANT ALL PRIVILEGES ON ALL TABLES IN SCHEMA public TO db_user;\nGRANT ALL PRIVILEGES ON ALL SEQUENCES IN SCHEMA public TO db_user;")

    add_h2("Step 3: Clone Codebase and Install Locked Dependencies")
    add_code_block("cd C:\\inetpub\\wwwroot\\\ngit clone https://[GIT_SERVER_URL]/DCMMS-with-Supabase.git DCMMS-App\ncd DCMMS-App\nnpm ci")

    add_h2("Step 4: Configure Production Environment File (.env)")
    add_code_block("DATABASE_URL=\"postgresql://db_user:StrongGovernmentPass2026!@127.0.0.1:5432/dmms_db?schema=public\"\nPORT=3000\nNODE_ENV=\"production\"\nNEXT_PUBLIC_APP_URL=\"http://localhost:3000\"")

    add_h2("Step 5: Synchronize Prisma Schema and Bootstrap System Admin")
    add_code_block("npx prisma generate\nnpx prisma db push\nnode scripts/create_system_admin.js")

    add_h2("Step 6: Build and Launch Production Cluster with PM2")
    add_code_block("npm run build\nnpm install -g pm2 pm2-windows-startup\npm2 start ecosystem.config.js\npm2 save\npm2-startup install")

    # ----------------------------------------------------
    # SECTION 7: CONFIGURATION MANAGEMENT
    # ----------------------------------------------------
    add_h1("7. Configuration Management & Environment Variables")
    add_custom_table(
        [Inches(2.0), Inches(1.0), Inches(1.2), Inches(2.3)],
        ["Variable Name", "Required", "Default", "Operational Purpose & Security Impact"],
        [
            ["DATABASE_URL", "YES", "postgresql://...", "Prisma & pg connection string. CRITICAL: Protect against credential leakage."],
            ["PORT", "YES", "3000", "Local listening port for Next.js HTTP server. LOW risk."],
            ["NODE_ENV", "YES", "production", "Toggles error detail suppression and performance caching. MEDIUM risk."],
            ["NEXT_PUBLIC_APP_URL", "YES", "http://localhost:3000", "Base URL for absolute hyperlink and document generation. LOW risk."]
        ]
    )

    # ----------------------------------------------------
    # SECTION 8: DATABASE ADMINISTRATION & SCHEMA REFERENCE
    # ----------------------------------------------------
    add_h1("8. Database Administration & Schema Reference")
    add_p("The DCMMS database contains 18 unified tables in PostgreSQL 16 under schema public:")

    add_h2("8.1 Summary Table Reference")
    add_custom_table(
        [Inches(2.2), Inches(1.3), Inches(3.0)],
        ["Table Name", "Primary Key", "Functional Purpose"],
        [
            ["register_officer_table", "id (Text / UUID)", "System users, security roles, and assigned subject categories."],
            ["daily_mail_letter_table", "id (BigInt Auto)", "Incoming complaint mail metadata and linked scanned PDF paths."],
            ["accused_officer_table", "id (UUID)", "Accused government officer records, NIC, designation, and school."],
            ["accused_school_table / institute_table", "id (BigInt Auto)", "Master registry of schools, zones, districts, and provinces."],
            ["subject_officer_form_table", "id (BigInt Auto)", "Master disciplinary case records, file numbers, and actions."],
            ["accused_officer_subject_officer_form_table", "Composite PK", "Many-to-many junction linking accused officers to case forms."],
            ["commitee_table", "id (UUID)", "Inquirer / committee member registry with NIC and active status."],
            ["school_table", "id (UUID)", "Inquirer affiliated schools and children's schools (Conflict check)."],
            ["chairment_by_case / members_by_case", "id (BigInt Auto)", "Case-level committee assignments for inquirers."],
            ["case_by_appointment_and_report_due_date", "id (UUID)", "Investigation appointment letter dates and statutory due dates."],
            ["case_by_date_extention", "id (BigInt Auto)", "Investigation deadline extension requests and approval status."],
            ["case_history", "history_id (UUID)", "Granular audit timeline tracking all case state modifications."],
            ["dcmms_audit_logs", "id (String)", "System security events: logins, toggles, password updates."],
            ["dcmms_sessions", "id (String)", "Active user session monitoring and remote forced termination."],
            ["dcmms_letter_edit_requests", "id (VarChar 100)", "Workflow for requesting corrections on registered daily mail."]
        ]
    )

    add_callout("High-Risk Database Alert", "Direct SQL modifications on subject_officer_form_table, altering constraints, or resetting autoincrement sequences must strictly be preceded by a verified pg_dump backup. Never run `npx prisma db push --force-reset` on production.", "CAUTION")

    # ----------------------------------------------------
    # SECTION 9: USER AND ROLE ADMINISTRATION
    # ----------------------------------------------------
    add_h1("9. User, Role, and Session Administration")
    add_p("DCMMS implements strict Role-Based Access Control (RBAC) across 5 system roles:")
    add_custom_table(
        [Inches(1.8), Inches(1.2), Inches(3.5)],
        ["System Role", "Route", "Functional Boundary & Permissions"],
        [
            ["System Administrator", "/system-admin", "Security logs, active session termination, user accounts, master Excel password."],
            ["Discipline Branch Admin", "/admin", "Case review, formal inquiry endorsements, extension request approvals, branch-wide reports."],
            ["Subject Officer", "/subject", "Case filing, minutes drafting, investigation requests, and tracking within assigned category."],
            ["Investigation Branch Admin", "/investigation", "Inquirer registry, committee appointment, conflict-of-interest validation, due dates."],
            ["Daily Mail Officer", "/daily-mail", "Incoming mail registration, PDF attachment, letter edit request submissions."]
        ]
    )

    # ----------------------------------------------------
    # SECTION 10: MASTER DATA & SYSTEM CONFIGURATION
    # ----------------------------------------------------
    add_h1("10. Master Data & System Configuration")
    add_p("DCMMS categorizes all national educational disciplinary cases across 15 master subject categories:")
    add_custom_table(
        [Inches(0.6), Inches(2.7), Inches(3.2)],
        ["Code", "Official Sinhala Name", "English Geographic / Operational Coverage"],
        [
            ["01", "ශාඛා ප්‍රධානී රාජකාරී", "Branch Head Duties & Direct Executive Matters"],
            ["02", "අධ්‍යාපන අමාත්‍යාංශය", "Ministry of Education Headquarters & Central Divisions"],
            ["03", "ශ්‍රී ජයවර්ධනපුර හා හෝමාගම කලාපය", "Sri Jayawardenepura & Homagama Educational Zone"],
            ["04", "කොළඹ කලාපය", "Colombo Educational Zone"],
            ["05", "ගම්පහ දිස්ත්‍රික්කය", "Gampaha District (All Zones)"],
            ["06", "කළුතර දිස්ත්‍රික්කය හා පිළියන්දල කලාපය", "Kalutara District & Piliyandala Zone"],
            ["07", "මාතර දිස්ත්‍රික්කය හා හම්බන්තොට දිස්ත්‍රික්කය", "Matara & Hambantota Districts (Southern Province)"],
            ["08", "මධ්‍යම පළාත", "Central Province (Kandy, Matale, Nuwara Eliya)"],
            ["09", "සබරගමුව පළාත", "Sabaragamuwa Province (Ratnapura, Kegalle)"],
            ["10", "උතුරු පළාත හා නැගෙනහිර පළාත", "Northern & Eastern Provinces"],
            ["11", "ගාල්ල දිස්ත්‍රික්කය", "Galle District"],
            ["12", "වයඹ පළාත", "North Western Province (Kurunegala, Puttalam)"],
            ["13", "විද්‍යාපීඨ", "National Colleges of Education (Vidya Peeta)"],
            ["14", "ඌව පළාත", "Uva Province (Badulla, Monaragala)"],
            ["15", "උතුරු මැද පළාත", "North Central Province (Anuradhapura, Polonnaruwa)"]
        ]
    )

    # ----------------------------------------------------
    # SECTION 11: CASE WORKFLOW
    # ----------------------------------------------------
    add_h1("11. End-to-End Case Lifecycle & Workflow Engine")
    add_p("The lifecycle of a disciplinary case follows a rigorous administrative trajectory:")
    add_bullet("1. Mail Registration: Daily Mail Officer enters letter metadata and attaches scanned PDF. Originating record created in daily_mail_letter_table.")
    add_bullet("2. Case Assignment & Filing: Subject Officer identifies matching subject code (01-15), links letter, enters accused officer details, and generates Master File Reference Number in subject_officer_form_table.")
    add_bullet("3. Inquiry Decision: Subject Officer determines whether case requires (A) Direct Warning/Closure, (B) Preliminary Investigation (PI), or (C) Formal Investigation (FI).")
    add_bullet("4. Inquirer Appointment: Investigation Branch Admin selects committee chair and members from commitee_table. Conflict-of-interest check automatically validates against school_table.")
    add_bullet("5. Appointment Letter & Due Date: Appointment Letter Date recorded; system calculates statutory Report Due Date (+30 or +60 days).")
    add_bullet("6. Extension Management: If inquirers require more time, an extension request is filed in case_by_date_extention and reviewed by Discipline Branch Admin.")
    add_bullet("7. Inquiry Report & Disciplinary Order: Inquirer report received and summarized; Discipline Branch Head endorses final Disciplinary Order; case marked Closed.")

    # ----------------------------------------------------
    # SECTION 12: BUSINESS RULES REPOSITORY
    # ----------------------------------------------------
    add_h1("12. Business Rules Repository")
    add_custom_table(
        [Inches(0.9), Inches(1.1), Inches(1.1), Inches(3.4)],
        ["Rule ID", "Module", "Target Role", "Business Rule & Enforcement Mechanism"],
        [
            ["BR-001", "Daily Mail", "Daily Mail", "Registered letters cannot be edited or deleted without formal edit request approved by Branch Admin."],
            ["BR-002", "Case Filing", "Subject Officer", "Subject Officers can only file and access cases within their assigned subject_type category."],
            ["BR-003", "Inquiry Setup", "Invest Admin", "Inquirer cannot be appointed if their affiliated schools match the accused officer's school (Conflict of Interest)."],
            ["BR-004", "Case Status", "All Users", "Closed cases become read-only. No extensions or committee modifications permitted without formal reopening."],
            ["BR-005", "File Upload", "Upload API", "Only valid PDF files (<= 25 MB) are accepted by the upload pipeline."],
            ["BR-006", "Excel Export", "All Users", "System-wide Excel report downloads require authentication against the master admin export password."],
            ["BR-007", "User Mgmt", "System Admin", "User accounts cannot be hard deleted; only toggled between Active and Deactivated."]
        ]
    )

    # ----------------------------------------------------
    # SECTION 13: NOTIFICATIONS & MONITORING
    # ----------------------------------------------------
    add_h1("13. Notification & Deadline Monitoring Engine")
    add_p("The system calculates statutory deadlines and generates dashboard alerts automatically:")
    add_bullet("Normal Status: Report Due Date is > 14 days in the future.")
    add_bullet("Warning Status: Report Due Date is within <= 14 days. Displayed as Amber Warning badge.")
    add_bullet("Overdue Alert: Report Due Date has elapsed without report submission. Displayed as Red Flashing badge and incremented in the Overdue Counter.")

    # ----------------------------------------------------
    # SECTION 14: DOCUMENT MANAGEMENT
    # ----------------------------------------------------
    add_h1("14. Document & File Storage Management")
    add_p("Uploaded scanned complaints and investigation reports are stored under public/uploads/documents/.")
    add_p("File Naming Algorithm: [Sanitized_Ref_No]_[Unix_Timestamp]_[Original_Filename].pdf\nExample: CR_2026_089_1725182400000_Complaint_Affidavit.pdf")
    add_p("NTFS Permissions: Ensure IIS_IUSRS or the PM2 service account has Read/Write/Modify permissions on the upload directory.")

    # ----------------------------------------------------
    # SECTION 15: AUDIT TRAIL AND LOGGING
    # ----------------------------------------------------
    add_h1("15. Audit Trail, Logging, and Compliance")
    add_p("DCMMS maintains dual audit streams:")
    add_bullet("Procedural Case Audit (case_history): Records all case status transitions, committee appointments, extensions, and remarks.")
    add_bullet("System Security Audit (dcmms_audit_logs & dcmms_sessions): Records logins, failed authentications, forced logouts, and account toggles.")

    # ----------------------------------------------------
    # SECTION 16: BACKUP AND RESTORE PROCEDURES
    # ----------------------------------------------------
    add_h1("16. Backup, Verification, and Restore Procedures")
    add_h2("16.1 Automated Backup Execution")
    add_code_block("powershell -ExecutionPolicy Bypass -File \"C:\\inetpub\\DCMMS-App\\scripts\\backup_postgres.ps1\" -DbUser \"db_user\" -DbName \"dmms_db\" -BackupDir \"C:\\Backups\\PostgreSQL\" -RetentionDays 30")

    add_h2("16.2 Database Restoration Procedure")
    add_code_block("# 1. Stop Next.js Application\npm2 stop dmms-app\n\n# 2. Terminate Database Connections & Re-create DB (in psql)\nSELECT pg_terminate_backend(pid) FROM pg_stat_activity WHERE datname = 'dmms_db' AND pid <> pg_backend_pid();\nDROP DATABASE dmms_db;\nCREATE DATABASE dmms_db;\n\n# 3. Restore Database from Dump\npg_restore -h 127.0.0.1 -p 5432 -U db_user -d dmms_db -v \"C:\\Backups\\PostgreSQL\\dmms_backup_[Timestamp].dump\"\n\n# 4. Restart Application\npm2 start dmms-app")

    # ----------------------------------------------------
    # SECTION 17: VERSION CONTROL
    # ----------------------------------------------------
    add_h1("17. Version Control & Branching Strategy")
    add_p("Maintain a standard Git flow with main (production releases), staging (UAT testing), and feature/ branches. All production releases must be tagged with Semantic Versioning (e.g. git tag -a v1.0.0 -m \"Production Release 1.0.0\").")

    # ----------------------------------------------------
    # SECTION 18 & 19: CHANGE & RELEASE MANAGEMENT
    # ----------------------------------------------------
    add_h1("18. Change & Release Management Lifecycle")
    add_p("Every system modification must proceed through formal stages: Change Request (CR) -> Impact Analysis -> Staging Testing & UAT -> Production Backup -> Deployment -> Verification Sign-off.")

    # ----------------------------------------------------
    # SECTION 20 & 21: DATABASE MIGRATION & ROLLBACK
    # ----------------------------------------------------
    add_h1("19. Database Migration & Emergency Rollback")
    add_h2("19.1 Prisma Schema Migration Protocol")
    add_code_block("# 1. Update schema.prisma\n# 2. Generate updated TypeScript client\nnpx prisma generate\n# 3. Apply schema to database\nnpx prisma db push")

    add_h2("19.2 Emergency Rollback Procedure")
    add_code_block("pm2 stop dmms-app\ngit checkout tags/v1.0.0\nnpm ci\npg_restore -h 127.0.0.1 -p 5432 -U db_user -d dmms_db -c -v \"C:\\Backups\\PostgreSQL\\dmms_pre_deploy_backup.dump\"\nnpm run build\npm2 restart dmms-app")

    # ----------------------------------------------------
    # SECTION 22 & 23: TESTING & TROUBLESHOOTING
    # ----------------------------------------------------
    add_h1("20. Testing Protocols & Troubleshooting Guide")
    add_custom_table(
        [Inches(1.8), Inches(1.8), Inches(2.9)],
        ["Issue Symptom", "Probable Root Cause", "Resolution Runbook"],
        [
            ["App crash / Port 3000 busy", "Orphaned Node.js process", "netstat -ano | findstr :3000; taskkill /F /PID [PID]; pm2 restart dmms-app"],
            ["Database connection refused", "PostgreSQL service stopped", "Start postgresql-x64-16 in services.msc; check disk space on C:\\"],
            ["PDF Upload fails (400/500)", "Non-PDF file or NTFS permissions", "Verify file is valid PDF <= 25MB; grant Full Control to IIS/PM2 user on public/uploads"],
            ["Excel Export password rejected", "Password modified in admin", "Log into /system-admin -> Excel Export Security widget to view/reset password"]
        ]
    )

    # ----------------------------------------------------
    # SECTION 24 - 28: SECURITY, PERFORMANCE & HOSTING
    # ----------------------------------------------------
    add_h1("21. Security, Performance & Infrastructure Administration")
    add_bullet("PostgreSQL Port 5432 must remain bound exclusively to 127.0.0.1 (Localhost only).")
    add_bullet("Application processes execute in PM2 Cluster Mode utilizing all available CPU cores.")
    add_bullet("Nightly automated database backups at 02:00 AM configured via Windows Task Scheduler.")
    add_bullet("Disaster Recovery RPO: 24 Hours (Nightly dump); RTO: 2 Hours (Bare-metal image restore).")

    # ----------------------------------------------------
    # SECTION 29 - 33: OPERATIONAL TASKS & MAINTENANCE
    # ----------------------------------------------------
    add_h1("22. Operational Tasks & Maintenance Checklists")
    add_h2("22.1 Daily Checklist (5 Minutes)")
    add_bullet("[ ] Verify PM2 status: pm2 status (Ensure dmms-app is online with 0 restarts).")
    add_bullet("[ ] Verify database backup log: C:\\Backups\\PostgreSQL\\backup_log.txt.")
    add_bullet("[ ] Check server drive free space (> 20% free space).")

    add_h2("22.2 Weekly Checklist (15 Minutes)")
    add_bullet("[ ] Review /system-admin security audit logs for unauthorized attempts.")
    add_bullet("[ ] Monitor active user sessions and force-logout stale connections.")
    add_bullet("[ ] Review upload directory storage growth in public/uploads/documents/.")

    add_h2("22.3 Monthly Checklist (1 Hour)")
    add_bullet("[ ] Test database restoration on development staging sandbox.")
    add_bullet("[ ] Deactivate transferred or retired staff accounts with Discipline Branch Head.")
    add_bullet("[ ] Transfer backups older than 30 days to secondary off-site government storage.")

    # ----------------------------------------------------
    # SECTION 34 - 38: IMPACT MATRIX, CHECKLISTS & APPENDICES
    # ----------------------------------------------------
    add_h1("23. System Change Impact Matrix")
    add_custom_table(
        [Inches(1.8), Inches(1.1), Inches(1.1), Inches(1.5), Inches(1.0)],
        ["Change Scenario", "DB Impact", "UI Impact", "Testing Required", "Risk"],
        [
            ["Add New Subject Category", "None", "Dropdowns", "Verify case creation in category", "LOW"],
            ["Add Accused Officer Field", "High (Migration)", "High (Forms)", "Full case filing smoke test", "MEDIUM"],
            ["Modify Due Date Logic", "Medium", "Medium", "Verify +30/+60 day calculations", "MEDIUM"],
            ["Update Inquirer Conflict Rule", "None", "High", "Run automated conflict test suite", "HIGH"],
            ["Next.js / React Upgrade", "None", "Critical", "Full regression test across all 5 roles", "HIGH"]
        ]
    )

    add_h1("24. Printable Production Release Sign-Off Checklist")
    add_p("================================================================================")
    add_p("DCMMS PRODUCTION RELEASE SIGN-OFF & VERIFICATION CHECKLIST")
    add_p("Release Version: _______________          Date of Release: _____________________")
    add_p("Release Engineer: ______________          Supervising Admin: ___________________")
    add_bullet("[ ] 1. Change Request formally approved by Discipline Branch Head.")
    add_bullet("[ ] 2. Code changes merged into main repository branch and tagged.")
    add_bullet("[ ] 3. Automated backup executed and dump file verified (Size > 0 KB).")
    add_bullet("[ ] 4. Target production server disk space verified (> 15 GB free).")
    add_bullet("[ ] 5. Dependencies installed cleanly via npm ci.")
    add_bullet("[ ] 6. Database schema push/migration executed cleanly via Prisma.")
    add_bullet("[ ] 7. Next.js production build succeeded (npm run build with 0 errors).")
    add_bullet("[ ] 8. PM2 application cluster reloaded (pm2 reload ecosystem.config.js).")
    add_bullet("[ ] 9. System Administrator login verified at /system-admin.")
    add_bullet("[ ] 10. Daily Mail letter registration and PDF upload verified.")
    add_bullet("[ ] 11. Subject Officer case view and update verified.")
    add_bullet("[ ] 12. Investigation Branch appointment workflow verified.")
    add_bullet("[ ] 13. Audit trail log verified recording the deployment event.")
    add_bullet("[ ] 14. Handover & Release Notes signed off by Lead Technical Officer.")
    add_p("SIGNATURE OF SYSTEM ADMINISTRATOR: __________________________ DATE: ____________")
    add_p("SIGNATURE OF DISCIPLINE BRANCH HEAD: ________________________ DATE: ____________")
    add_p("================================================================================")

    add_h1("25. Appendices")
    add_h2("Appendix A: Emergency Contact & Escalation Roster")
    add_custom_table(
        [Inches(2.2), Inches(2.1), Inches(2.2)],
        ["Contact Role", "Designated Officer Name", "Contact Details"],
        [
            ["Lead System Administrator", "[TO BE COMPLETED BY FUTURE SYSADMIN]", "[TO BE CONFIRMED]"],
            ["Database Administrator (DBA)", "[TO BE COMPLETED BY FUTURE DBA]", "[TO BE CONFIRMED]"],
            ["Discipline Branch Head", "[TO BE CONFIRMED - Senior Asst. Secretary]", "[TO BE CONFIRMED]"],
            ["Director of ICT (MOE)", "[TO BE CONFIRMED - Director of ICT]", "[TO BE CONFIRMED]"]
        ]
    )

    output_path = os.path.join(os.getcwd(), "DCMMS_System_Administrator_Handover_Guide.docx")
    doc.save(output_path)
    print(f"Successfully generated: {output_path}")

if __name__ == "__main__":
    create_full_document()
