import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)

def set_table_borders(table, color="D3D3D3"):
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        borders = parse_xml(f'''
            <w:tblBorders {nsdecls("w")}>
                <w:top w:val="single" w:sz="4" w:space="0" w:color="{color}"/>
                <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{color}"/>
                <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{color}"/>
                <w:insideV w:val="none"/>
                <w:left w:val="none"/>
                <w:right w:val="none"/>
            </w:tblBorders>
        ''')
        tblPr[0].append(borders)

def add_callout(doc, title, text, box_type="NOTE"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    
    color_map = {
        "NOTE": ("EBF8FF", "2B6CB0"),
        "IMPORTANT": ("E6FFFA", "234E52"),
        "WARNING": ("FFFAF0", "C05621"),
        "CAUTION": ("FFF5F5", "C53030")
    }
    bg_hex, border_hex = color_map.get(box_type, ("F7FAFC", "4A5568"))
    
    cell = tbl.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, bg_hex)
    set_cell_margins(cell, top=140, bottom=140, left=200, right=180)
    
    tcPr = cell._element.get_or_add_tcPr()
    borders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:left w:val="single" w:sz="24" w:space="0" w:color="{border_hex}"/>
            <w:top w:val="none"/>
            <w:right w:val="none"/>
            <w:bottom w:val="none"/>
        </w:tcBorders>
    ''')
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run_title = p.add_run(f"[{box_type}] {title}\n")
    run_title.bold = True
    run_title.font.name = "Arial"
    run_title.font.size = Pt(10)
    r, g, b = tuple(int(border_hex[i:i+2], 16) for i in (0, 2, 4))
    run_title.font.color.rgb = RGBColor(r, g, b)
    
    run_text = p.add_run(text)
    run_text.font.name = "Calibri"
    run_text.font.size = Pt(9.5)
    run_text.font.color.rgb = RGBColor(45, 55, 72)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def add_code_block(doc, code_text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    
    cell = tbl.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, "F8F9FA")
    set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
    
    tcPr = cell._element.get_or_add_tcPr()
    borders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:left w:val="single" w:sz="12" w:space="0" w:color="0A2540"/>
            <w:top w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>
            <w:right w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>
            <w:bottom w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>
        </w:tcBorders>
    ''')
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(code_text.strip())
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(30, 41, 59)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def format_table(table, col_widths, headers, data):
    set_table_borders(table)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header row
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "0A2540")
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=120, right=120)
        p = hdr_cells[i].paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.bold = True
            run.font.name = "Arial"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(255, 255, 255)
            
    # Data rows
    for row_idx, row_data in enumerate(data):
        row_cells = table.add_row().cells
        bg_color = "F8FAFC" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, text_val in enumerate(row_data):
            row_cells[col_idx].text = str(text_val)
            set_cell_background(row_cells[col_idx], bg_color)
            set_cell_margins(row_cells[col_idx], top=80, bottom=80, left=100, right=100)
            p = row_cells[col_idx].paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            for run in p.runs:
                run.font.name = "Calibri"
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(30, 41, 59)
                
    # Widths
    for row in table.rows:
        for idx, width in enumerate(col_widths):
            row.cells[idx].width = width

print("Helper definitions complete.")
